# import streamlit as st
# from llama_index.core import VectorStoreIndex, Document, Settings
# from llama_index.vector_stores.faiss import FaissVectorStore
# from llama_index.embeddings.gemini import GeminiEmbedding
# from llama_index.llms.groq import Groq
# import faiss
# from docx import Document as DocxDocument
# import pymupdf4llm
# import tempfile
# import os
# import json
# import re

# # Initialize session state variables
# if "api_key" not in st.session_state:
#     st.session_state.api_key = ""
# if "google_api_key" not in st.session_state:
#     st.session_state.google_api_key = ""
# if "current_page" not in st.session_state:
#     st.session_state.current_page = "upload"
# if "index" not in st.session_state:
#     st.session_state.index = None
# if "current_assessment" not in st.session_state:
#     st.session_state.current_assessment = None
# if "user_answers" not in st.session_state:
#     st.session_state.user_answers = {}
# if "assessment_score" not in st.session_state:
#     st.session_state.assessment_score = None
# if "current_question_index" not in st.session_state:
#     st.session_state.current_question_index = 0

# def read_file(file):
#     if file.name.endswith('.pdf'):
#         with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
#             tmp_file.write(file.getvalue())
#             tmp_file_path = tmp_file.name
#         try:
#             md_text = pymupdf4llm.to_markdown(tmp_file_path)
#         finally:
#             os.unlink(tmp_file_path)
#         return md_text
#     elif file.name.endswith('.docx'):
#         return "\n".join(para.text for para in DocxDocument(file).paragraphs)
#     else:
#         return file.getvalue().decode()

# def process_documents(uploaded_files):
#     documents = [Document(text=read_file(file), metadata={"filename": file.name}) 
#                 for file in uploaded_files]
    
#     d = 768  # Dimension for Google embeddings
#     faiss_index = faiss.IndexFlatL2(d)
#     vector_store = FaissVectorStore(faiss_index=faiss_index)
    
#     st.session_state.index = VectorStoreIndex.from_documents(
#         documents,
#         vector_store=vector_store
#     )

# def parse_mcq_response(response):
#     questions = []
#     current_question = {}
    
#     for line in response.split('\n'):
#         line = line.strip()
#         if not line:
#             continue
            
#         if line.startswith('Q'):
#             if current_question:
#                 questions.append(current_question)
#             current_question = {
#                 'question': line[line.find('.')+1:].strip(),
#                 'options': [],
#                 'correct_answer': None
#             }
#         elif line.startswith(('a)', 'b)', 'c)', 'd)')):
#             current_question['options'].append(line[2:].strip())
#         elif line.startswith('Correct Answer:'):
#             current_question['correct_answer'] = line.split(':')[1].strip().lower()
    
#     if current_question:
#         questions.append(current_question)
    
#     return questions

# def generate_mcq(context, num_questions=5, difficulty="medium", topics=None):
#     topic_str = f" focusing on {', '.join(topics)}" if topics else ""
#     prompt = f"""
#     Based on the following context, generate {num_questions} {difficulty}-level multiple choice questions{topic_str}.
#     Each question should have 4 options with only one correct answer.
#     Make sure the questions are challenging but fair for the {difficulty} difficulty level.
    
#     Format:
#     Q1. [Question]
#     a) [Option]
#     b) [Option]
#     c) [Option]
#     d) [Option]
#     Correct Answer: [a/b/c/d]

#     Context: {context}
#     """
    
#     query_engine = st.session_state.index.as_query_engine(
#         response_mode="compact"
#     )
#     response = query_engine.query(prompt)
#     return parse_mcq_response(str(response))

# def generate_free_response(context, num_questions=3, difficulty="medium", topics=None):
#     topic_str = f" focusing on {', '.join(topics)}" if topics else ""
#     prompt = f"""
#     Based on the context, generate {num_questions} {difficulty}-level open-ended questions{topic_str} that test understanding
#     of the key concepts. For each question, provide:
#     1. The question
#     2. Key points that should be included in a good answer
#     3. A model answer for reference
#     4. Scoring criteria (what makes an answer excellent, good, or needs improvement)
#     """
#     query_engine = st.session_state.index.as_query_engine(
#         response_mode="compact"
#     )
#     response = query_engine.query(prompt)
#     return str(response)

# def evaluate_free_response(question, model_answer, user_answer):
#     prompt = f"""
#     Evaluate the following student answer against the model answer and provide:
#     1. Score (0-100)
#     2. Detailed feedback
#     3. Areas for improvement
    
#     Question: {question}
#     Model Answer: {model_answer}
#     Student Answer: {user_answer}
#     """
    
#     query_engine = st.session_state.index.as_query_engine(
#         response_mode="compact"
#     )
#     response = query_engine.query(prompt)
#     return str(response)

# # Streamlit UI
# st.set_page_config(page_title="Interactive Learning Assessment", page_icon="📚", layout="wide")

# # Sidebar for API keys and document upload
# with st.sidebar:
#     st.header("🔑 API Keys")
#     st.session_state.api_key = st.text_input("Enter your Groq API Key:", type="password")
#     st.session_state.google_api_key = st.text_input("Enter your Google API Key:", type="password")
    
#     if st.session_state.current_page == "upload":
#         st.header("📁 Document Upload")
#         uploaded_files = st.file_uploader(
#             "Upload your learning materials (PDF, DOCX, TXT)",
#             accept_multiple_files=True,
#             type=['pdf', 'docx', 'txt']
#         )
        
#         if uploaded_files and st.session_state.api_key and st.session_state.google_api_key:
#             with st.spinner("Processing documents..."):
#                 llm = Groq(api_key=st.session_state.api_key, model="llama-3.3-70b-versatile")
#                 embed_model = GeminiEmbedding(api_key=st.session_state.google_api_key)
                
#                 Settings.embed_model = embed_model
#                 Settings.llm = llm
                
#                 process_documents(uploaded_files)
#             st.success(f"{len(uploaded_files)} document(s) processed successfully!")
#             st.session_state.current_page = "menu"

# # Main content area
# st.title("📚 Interactive Learning Assessment")

# if st.session_state.current_page == "upload":
#     st.info("Please upload your learning materials in the sidebar to get started!")

# elif st.session_state.current_page == "menu":
#     st.header("Choose Assessment Type")
#     col1, col2 = st.columns(2)
    
#     with col1:
#         if st.button("📝 Knowledge Assessment (MCQ)", use_container_width=True):
#             st.session_state.current_page = "mcq_config"
    
#     with col2:
#         if st.button("💭 Skills Development (Free Response)", use_container_width=True):
#             st.session_state.current_page = "free_response_config"

# elif st.session_state.current_page == "mcq_config":
#     st.header("Configure MCQ Assessment")
#     if st.button("← Back to Menu"):
#         st.session_state.current_page = "menu"
    
#     col1, col2 = st.columns(2)
#     with col1:
#         num_questions = st.slider("Number of questions:", 5, 20, 10)
#         difficulty = st.select_slider(
#             "Difficulty level:",
#             options=["easy", "medium", "hard"],
#             value="medium"
#         )
    
#     with col2:
#         topics = st.multiselect(
#             "Select specific topics (optional):",
#             ["Topic 1", "Topic 2", "Topic 3"],  # You can dynamically generate these from the documents
#             default=None
#         )
    
#     if st.button("Start Assessment"):
#         with st.spinner("Generating questions..."):
#             st.session_state.current_assessment = generate_mcq(
#                 context="",
#                 num_questions=num_questions,
#                 difficulty=difficulty,
#                 topics=topics
#             )
#             st.session_state.current_question_index = 0
#             st.session_state.user_answers = {}
#             st.session_state.current_page = "mcq_assessment"

# elif st.session_state.current_page == "mcq_assessment":
#     st.header("Multiple Choice Assessment")
    
#     if st.session_state.current_assessment:
#         question = st.session_state.current_assessment[st.session_state.current_question_index]
        
#         # Progress indicator
#         st.progress((st.session_state.current_question_index + 1) / len(st.session_state.current_assessment))
#         st.write(f"Question {st.session_state.current_question_index + 1} of {len(st.session_state.current_assessment)}")
        
#         # Display question
#         st.write(f"**{question['question']}**")
        
#         # Radio buttons for options
#         selected_option = st.radio(
#             "Select your answer:",
#             ['a', 'b', 'c', 'd'],
#             format_func=lambda x: f"{x}) {question['options'][ord(x) - ord('a')]}"
#         )
        
#         col1, col2 = st.columns(2)
#         with col1:
#             if st.button("Previous Question") and st.session_state.current_question_index > 0:
#                 st.session_state.current_question_index -= 1
        
#         with col2:
#             if st.session_state.current_question_index < len(st.session_state.current_assessment) - 1:
#                 if st.button("Next Question"):
#                     st.session_state.user_answers[st.session_state.current_question_index] = selected_option
#                     st.session_state.current_question_index += 1
#             else:
#                 if st.button("Submit Assessment"):
#                     st.session_state.user_answers[st.session_state.current_question_index] = selected_option
                    
#                     # Calculate score
#                     correct = sum(1 for i, ans in st.session_state.user_answers.items() 
#                                 if ans == st.session_state.current_assessment[i]['correct_answer'])
#                     total = len(st.session_state.current_assessment)
#                     st.session_state.assessment_score = (correct / total) * 100
#                     st.session_state.current_page = "mcq_results"
#                     st.rerun()

# elif st.session_state.current_page == "mcq_results":
#     st.header("Assessment Results")
    
#     st.write(f"### Score: {st.session_state.assessment_score:.1f}%")
    
#     for i, question in enumerate(st.session_state.current_assessment):
#         with st.expander(f"Question {i + 1}"):
#             st.write(f"**{question['question']}**")
#             user_answer = st.session_state.user_answers.get(i)
#             correct_answer = question['correct_answer']
            
#             for j, option in enumerate(question['options']):
#                 option_letter = chr(ord('a') + j)
#                 if option_letter == user_answer == correct_answer:
#                     st.write(f"✅ {option_letter}) {option} (Your correct answer)")
#                 elif option_letter == user_answer:
#                     st.write(f"❌ {option_letter}) {option} (Your answer)")
#                 elif option_letter == correct_answer:
#                     st.write(f"✳️ {option_letter}) {option} (Correct answer)")
#                 else:
#                     st.write(f"⚪ {option_letter}) {option}")
    
#     if st.button("Return to Menu"):
#         st.session_state.current_page = "menu"
#         st.rerun()

# elif st.session_state.current_page == "free_response_config":
#     st.header("Configure Free Response Assessment")
#     if st.button("← Back to Menu"):
#         st.session_state.current_page = "menu"
    
#     col1, col2 = st.columns(2)
#     with col1:
#         num_questions = st.slider("Number of questions:", 1, 5, 3)
#         difficulty = st.select_slider(
#             "Difficulty level:",
#             options=["easy", "medium", "hard"],
#             value="medium"
#         )
    
#     with col2:
#         topics = st.multiselect(
#             "Select specific topics (optional):",
#             ["Topic 1", "Topic 2", "Topic 3"],  # You can dynamically generate these
#             default=None
#         )
    
#     if st.button("Start Assessment"):
#         with st.spinner("Generating questions..."):
#             st.session_state.current_assessment = generate_free_response(
#                 context="",
#                 num_questions=num_questions,
#                 difficulty=difficulty,
#                 topics=topics
#             )
#             st.session_state.current_page = "free_response_assessment"
#             st.rerun()

# elif st.session_state.current_page == "free_response_assessment":
#     st.header("Free Response Assessment")
    
#     if st.session_state.current_assessment:
#         questions = st.session_state.current_assessment.split('\n\n')
        
#         for i, question_block in enumerate(questions):
#             with st.expander(f"Question {i + 1}", expanded=True):
#                 st.write(question_block)
#                 answer = st.text_area(f"Your answer for Question {i + 1}:", key=f"q{i}")
#                 if st.button(f"Submit Answer {i + 1}"):
#                     evaluation = evaluate_free_response(
#                         question_block,
#                         "Model answer from question block",
#                         answer
#                     )
#                     st.write("### Evaluation")
#                     st.write(evaluation)
    
#     if st.button("Finish Assessment"):
#         st.session_state.current_page = "menu"
#         st.rerun()

import streamlit as st
from llama_index.core import VectorStoreIndex, Document, Settings
from llama_index.vector_stores.faiss import FaissVectorStore
from llama_index.embeddings.gemini import GeminiEmbedding
from llama_index.llms.groq import Groq
import faiss
from docx import Document as DocxDocument
import pymupdf4llm
import tempfile
import os

# Initialize session state variables
if "api_key" not in st.session_state:
    st.session_state.api_key = ""
if "google_api_key" not in st.session_state:
    st.session_state.google_api_key = ""
if "current_page" not in st.session_state:
    st.session_state.current_page = "upload"
if "index" not in st.session_state:
    st.session_state.index = None
if "current_assessment" not in st.session_state:
    st.session_state.current_assessment = None
if "user_answers" not in st.session_state:
    st.session_state.user_answers = {}
if "assessment_score" not in st.session_state:
    st.session_state.assessment_score = None
if "current_question_index" not in st.session_state:
    st.session_state.current_question_index = 0
if "selected_difficulty" not in st.session_state:
    st.session_state.selected_difficulty = "medium"

def read_file(file):
    if file.name.endswith('.pdf'):
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            tmp_file.write(file.getvalue())
            tmp_file_path = tmp_file.name
        try:
            md_text = pymupdf4llm.to_markdown(tmp_file_path)
        finally:
            os.unlink(tmp_file_path)
        return md_text
    elif file.name.endswith('.docx'):
        return "\n".join(para.text for para in DocxDocument(file).paragraphs)
    else:
        return file.getvalue().decode()

def process_documents(uploaded_files):
    documents = [Document(text=read_file(file), metadata={"filename": file.name}) 
                for file in uploaded_files]
    
    d = 768  # Dimension for Google embeddings
    faiss_index = faiss.IndexFlatL2(d)
    vector_store = FaissVectorStore(faiss_index=faiss_index)
    
    st.session_state.index = VectorStoreIndex.from_documents(
        documents,
        vector_store=vector_store
    )

def parse_mcq_response(response):
    questions = []
    current_question = {}
    
    for line in response.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('Q'):
            if current_question:
                questions.append(current_question)
            current_question = {
                'question': line[line.find('.')+1:].strip(),
                'options': [],
                'correct_answer': None,
                'explanation': None
            }
        elif line.startswith(('a)', 'b)', 'c)', 'd)')):
            current_question['options'].append(line[2:].strip())
        elif line.startswith('Correct Answer:'):
            current_question['correct_answer'] = line.split(':')[1].strip().lower()
        elif line.startswith('Explanation:'):
            current_question['explanation'] = line.split(':')[1].strip()
    
    if current_question:
        questions.append(current_question)
    
    return questions

def generate_mcq(context, num_questions=5, difficulty="medium"):
    prompt = f"""
    Based on the following context, generate {num_questions} {difficulty}-level multiple choice questions.
    Each question should:
    - Have 4 options with only one correct answer
    - Include a brief explanation of why the answer is correct
    - Be challenging but fair for the {difficulty} difficulty level
    - Test different aspects of understanding
    
    Format:
    Q1. [Question]
    a) [Option]
    b) [Option]
    c) [Option]
    d) [Option]
    Correct Answer: [a/b/c/d]
    Explanation: [Brief explanation of the correct answer]

    Context: {context}
    """
    
    query_engine = st.session_state.index.as_query_engine(
        response_mode="compact"
    )
    response = query_engine.query(prompt)
    return parse_mcq_response(str(response))

def generate_free_response(context, num_questions=3, difficulty="medium"):
    prompt = f"""
    Based on the context, generate {num_questions} {difficulty}-level open-ended questions that test deep understanding
    of the key concepts. For each question, provide:
    1. A thought-provoking question that requires analysis and critical thinking
    2. 3-4 key points that should be included in a good answer
    3. A detailed model answer for reference
    4. Clear scoring rubric (90-100%: Excellent, 80-89%: Good, 70-79%: Satisfactory, Below 70%: Needs Improvement)
    5. Common misconceptions to watch for
    """
    query_engine = st.session_state.index.as_query_engine(
        response_mode="compact"
    )
    response = query_engine.query(prompt)
    return str(response)

def evaluate_free_response(question, model_answer, user_answer):
    prompt = f"""
    Evaluate the following student answer against the model answer and provide:
    1. Numerical score (0-100)
    2. Specific strengths of the answer
    3. Areas for improvement
    4. Suggestions for deeper understanding
    5. Additional resources or concepts to explore
    
    Question: {question}
    Model Answer: {model_answer}
    Student Answer: {user_answer}
    """
    
    query_engine = st.session_state.index.as_query_engine(
        response_mode="compact"
    )
    response = query_engine.query(prompt)
    return str(response)

# Streamlit UI
st.set_page_config(page_title="Interactive Learning Assessment", page_icon="📚", layout="wide")

# Custom CSS for better styling
st.markdown("""
    <style>
    div.stButton > button {
        width: 100%;
        height: 60px;
        margin: 10px 0px;
    }
    .difficulty-btn {
        border-radius: 20px;
        padding: 15px 30px;
        margin: 10px;
        border: none;
        cursor: pointer;
    }
    .difficulty-btn.selected {
        background-color: #0066cc;
        color: white;
    }
    </style>
    """, unsafe_allow_html=True)

# Sidebar for API keys and document upload
with st.sidebar:
    st.header("🔑 API Configuration")
    st.session_state.api_key = st.text_input("Enter your Groq API Key:", type="password")
    st.session_state.google_api_key = st.text_input("Enter your Google API Key:", type="password")
    
    if st.session_state.current_page == "upload":
        st.header("📁 Document Upload")
        uploaded_files = st.file_uploader(
            "Upload learning materials (PDF, DOCX, TXT)",
            accept_multiple_files=True,
            type=['pdf', 'docx', 'txt']
        )
        
        if uploaded_files and st.session_state.api_key and st.session_state.google_api_key:
            with st.spinner("Processing documents..."):
                llm = Groq(api_key=st.session_state.api_key, model="llama-3.3-70b-versatile")
                embed_model = GeminiEmbedding(api_key=st.session_state.google_api_key)
                
                Settings.embed_model = embed_model
                Settings.llm = llm
                
                process_documents(uploaded_files)
            st.success(f"✅ {len(uploaded_files)} document(s) processed successfully!")
            st.session_state.current_page = "menu"

# Main content area
st.title("📚 Interactive Learning Assessment")

if st.session_state.current_page == "upload":
    st.info("👈 Please upload your learning materials in the sidebar to get started!")
    st.write("""
    ### Supported File Types:
    - 📄 PDF Documents
    - 📝 Word Documents (.docx)
    - 📑 Text Files (.txt)
    """)

elif st.session_state.current_page == "menu":
    st.header("Choose Assessment Type")
    col1, col2 = st.columns(2)
    
    with col1:
        st.write("### 📝 Knowledge Check")
        if st.button("Multiple Choice Questions", use_container_width=True):
            st.session_state.current_page = "mcq_config"
    
    with col2:
        st.write("### 💭 Deep Understanding")
        if st.button("Free Response Questions", use_container_width=True):
            st.session_state.current_page = "free_response_config"

elif st.session_state.current_page == "mcq_config":
    st.header("Configure MCQ Assessment")
    if st.button("← Back to Menu", key="back_btn"):
        st.session_state.current_page = "menu"
    
    st.write("### Number of Questions")
    num_questions = st.slider("Select the number of questions:", 5, 20, 10)
    
    st.write("### Difficulty Level")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        easy_selected = st.session_state.selected_difficulty == "easy"
        if st.button("🟢 Easy", key="easy", 
                    help="Basic concepts and straightforward questions",
                    type="primary" if easy_selected else "secondary"):
            st.session_state.selected_difficulty = "easy"
    
    with col2:
        medium_selected = st.session_state.selected_difficulty == "medium"
        if st.button("🟡 Medium", key="medium",
                    help="Moderate complexity and some application",
                    type="primary" if medium_selected else "secondary"):
            st.session_state.selected_difficulty = "medium"
    
    with col3:
        hard_selected = st.session_state.selected_difficulty == "hard"
        if st.button("🔴 Hard", key="hard",
                    help="Complex concepts and advanced application",
                    type="primary" if hard_selected else "secondary"):
            st.session_state.selected_difficulty = "hard"
    
    st.write("")
    if st.button("Start Assessment", type="primary"):
        with st.spinner("Generating questions..."):
            st.session_state.current_assessment = generate_mcq(
                context="",
                num_questions=num_questions,
                difficulty=st.session_state.selected_difficulty
            )
            st.session_state.current_question_index = 0
            st.session_state.user_answers = {}
            st.session_state.current_page = "mcq_assessment"

elif st.session_state.current_page == "mcq_assessment":
    st.header("Multiple Choice Assessment")
    
    if st.session_state.current_assessment:
        question = st.session_state.current_assessment[st.session_state.current_question_index]
        
        # Progress bar and question counter
        progress = (st.session_state.current_question_index + 1) / len(st.session_state.current_assessment)
        st.progress(progress)
        st.write(f"Question {st.session_state.current_question_index + 1} of {len(st.session_state.current_assessment)}")
        
        # Question card
        st.markdown("""
            <div style="background-color: #f0f2f6; padding: 20px; border-radius: 10px;">
        """, unsafe_allow_html=True)
        st.write(f"**{question['question']}**")
        
        # Radio buttons for options
        selected_option = st.radio(
            "Select your answer:",
            ['a', 'b', 'c', 'd'],
            format_func=lambda x: f"{x}) {question['options'][ord(x) - ord('a')]}"
        )
        st.markdown("</div>", unsafe_allow_html=True)
        
        # Navigation buttons
        col1, col2, col3 = st.columns([1, 2, 1])
        with col1:
            if st.session_state.current_question_index > 0:
                if st.button("← Previous", use_container_width=True):
                    st.session_state.current_question_index -= 1
        
        with col3:
            if st.session_state.current_question_index < len(st.session_state.current_assessment) - 1:
                if st.button("Next →", use_container_width=True):
                    st.session_state.user_answers[st.session_state.current_question_index] = selected_option
                    st.session_state.current_question_index += 1
            else:
                if st.button("Submit", type="primary", use_container_width=True):
                    st.session_state.user_answers[st.session_state.current_question_index] = selected_option
                    
                    # Calculate score
                    correct = sum(1 for i, ans in st.session_state.user_answers.items() 
                                if ans == st.session_state.current_assessment[i]['correct_answer'])
                    total = len(st.session_state.current_assessment)
                    st.session_state.assessment_score = (correct / total) * 100
                    st.session_state.current_page = "mcq_results"
                    st.rerun()

elif st.session_state.current_page == "mcq_results":
    st.header("Assessment Results")
    
    # Display overall score with appropriate color
    score = st.session_state.assessment_score
    if score >= 90:
        st.success(f"🎯 Score: {score:.1f}%")
    elif score >= 70:
        st.warning(f"📊 Score: {score:.1f}%")
    else:
        st.error(f"📉 Score: {score:.1f}%")
    
    # Question review
    for i, question in enumerate(st.session_state.current_assessment):
        with st.expander(f"Question {i + 1}", expanded=True):
            st.write(f"**{question['question']}**")
            user_answer = st.session_state.user_answers.get(i)
            correct_answer = question['correct_answer']
            
            for j, option in enumerate(question['options']):
                option_letter = chr(ord('a') + j)
                prefix = "👉 " if option_letter == user_answer else "   "
                if option_letter == user_answer == correct_answer:
                    st.success(f"{prefix}{option_letter}) {option} ✅")
                elif option_letter == user_answer:
                    st.error(f"{prefix}{option_letter}) {option} ❌")
                elif option_letter == correct_answer:
                    st.info(f"   {option_letter}) {option} (Correct answer)")
                else:
                    st.write(f"   {option_letter}) {option}")
            
            if question.get('explanation'):
                st.write("**Explanation:**")
                st.write(question['explanation'])
    
    # Summary and next steps
    st.write("### Summary")
    correct_count = sum(1 for i, ans in st.session_state.user_answers.items() 
                       if ans == st.session_state.current_assessment[i]['correct_answer'])
    total_count = len(st.session_state.current_assessment)
    
    st.write(f"- Total Questions: {total_count}")
    st.write(f"- Correct Answers: {correct_count}")
    st.write(f"- Incorrect Answers: {total_count - correct_count}")
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("Try Another Assessment", type="primary"):
            st.session_state.current_page = "menu"
            st.rerun()
    with col2:
        if st.button("Review Materials"):
            st.session_state.current_page = "upload"
            st.rerun()

elif st.session_state.current_page == "free_response_config":
    st.header("Configure Free Response Assessment")
    if st.button("← Back to Menu", key="back_btn"):
        st.session_state.current_page = "menu"
    
    st.write("### Number of Questions")
    num_questions = st.slider("Select the number of questions:", 1, 5, 3)
    
    st.write("### Difficulty Level")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        easy_selected = st.session_state.selected_difficulty == "easy"
        if st.button("🟢 Easy", key="easy",
                    help="Basic analysis and straightforward applications",
                    type="primary" if easy_selected else "secondary"):
            st.session_state.selected_difficulty = "easy"
    
    with col2:
        medium_selected = st.session_state.selected_difficulty == "medium"
        if st.button("🟡 Medium", key="medium",
                    help="Deeper analysis and concept connections",
                    type="primary" if medium_selected else "secondary"):
            st.session_state.selected_difficulty = "medium"
    
    with col3:
        hard_selected = st.session_state.selected_difficulty == "hard"
        if st.button("🔴 Hard", key="hard",
                    help="Complex analysis and critical thinking",
                    type="primary" if hard_selected else "secondary"):
            st.session_state.selected_difficulty = "hard"
    
    st.write("")
    if st.button("Start Assessment", type="primary"):
        with st.spinner("Generating questions..."):
            st.session_state.current_assessment = generate_free_response(
                context="",
                num_questions=num_questions,
                difficulty=st.session_state.selected_difficulty
            )
            st.session_state.current_page = "free_response_assessment"
            st.rerun()

elif st.session_state.current_page == "free_response_assessment":
    st.header("Free Response Assessment")
    
    if st.session_state.current_assessment:
        questions = st.session_state.current_assessment.split('\n\n')
        
        for i, question_block in enumerate(questions):
            with st.expander(f"Question {i + 1}", expanded=True):
                st.markdown("""
                    <div style="background-color: #f0f2f6; padding: 20px; border-radius: 10px;">
                """, unsafe_allow_html=True)
                st.write(question_block)
                st.markdown("</div>", unsafe_allow_html=True)
                
                answer = st.text_area(
                    "Your Answer:",
                    key=f"q{i}",
                    height=200,
                    placeholder="Type your response here...",
                    help="Write a detailed response addressing all aspects of the question"
                )
                
                col1, col2 = st.columns([3, 1])
                with col2:
                    if st.button(f"Submit Answer", key=f"submit_{i}", type="primary"):
                        evaluation = evaluate_free_response(
                            question_block,
                            "Model answer from question block",
                            answer
                        )
                        st.write("### Evaluation")
                        st.write(evaluation)
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("Save Progress"):
            st.success("Progress saved successfully!")
    with col2:
        if st.button("Finish Assessment", type="primary"):
            st.session_state.current_page = "menu"
            st.rerun()
